# app.py

import streamlit as st
import google.generativeai as genai
import google.api_core.exceptions
import os
import tempfile
import time
from dotenv import load_dotenv
import re
import json
from docx import Document
import io
import math
from datetime import datetime, timedelta
import pandas as pd
import plotly.express as px

# Import pydub for audio processing
from pydub import AudioSegment

# --- Database & Authentication Imports ---
from database import SessionLocal, init_db
from models import User, Transcript

# --- 1. INITIAL APP CONFIGURATION ---
st.set_page_config(
    page_title="Pro Bengali Transcriber",
    page_icon=" workbench",
    layout="wide"
)

# --- Initialize the database (creates tables if they don't exist) ---
init_db()

# --- 2. LOAD SECRETS & CONFIGURE API ---
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    st.error("FATAL: GEMINI_API_KEY is not configured on the server!")
    st.stop()
try:
    genai.configure(api_key=GEMINI_API_KEY)
except Exception as e:
    st.error(f"Failed to configure Google Gemini API: {e}")
    st.stop()

# --- 3. AUTHENTICATION & DATABASE HELPERS ---
def authenticate_user(username, password):
    db = SessionLocal()
    try:
        user = db.query(User).filter(User.username == username).first()
        if user and user.check_password(password):
            return user
        return None
    finally:
        db.close()

def add_new_user(username, password, role):
    db = SessionLocal()
    try:
        if db.query(User).filter(User.username == username).first():
            return False, "User already exists."
        new_user = User(username=username, role=role)
        new_user.set_password(password)
        db.add(new_user)
        db.commit()
        return True, f"User '{username}' created."
    finally:
        db.close()

def get_all_users_from_db():
    db = SessionLocal()
    try:
        return db.query(User.username, User.role).all()
    finally:
        db.close()

def save_transcript_to_db(user_id, title, filename, content):
    db = SessionLocal()
    try:
        new_transcript = Transcript(
            title=title,
            original_filename=filename,
            content=content,
            owner_id=user_id
        )
        db.add(new_transcript)
        db.commit()
        return True
    except Exception as e:
        st.error(f"Database Error: {e}")
        return False
    finally:
        db.close()

def get_user_transcripts(user_id):
    db = SessionLocal()
    try:
        return db.query(Transcript).filter(Transcript.owner_id == user_id).order_by(Transcript.created_at.desc()).all()
    finally:
        db.close()

def load_transcript_from_db(transcript_id):
    db = SessionLocal()
    try:
        return db.query(Transcript).filter(Transcript.id == transcript_id).first()
    finally:
        db.close()

# --- 4. CORE TRANSCRIPTION & AI FUNCTIONS ---
@st.cache_data(show_spinner="Translating transcript to English...", persist=True)
def _translate_text_with_gemini(transcript_data):
    model = genai.GenerativeModel(model_name="models/gemini-1.5-flash-latest")
    texts_to_translate = [segment['text'] for segment in transcript_data]
    prompt = (
        f"You are an expert Bengali to English translator. Translate each of the following Bengali text segments into English. "
        f"Your response MUST be a valid JSON array of strings, with each string being the translation of the corresponding input text. "
        f"Maintain the original order and number of segments. Do not add any commentary.\n\n"
        f"**JSON Array to Translate:**\n{json.dumps(texts_to_translate, ensure_ascii=False)}"
    )
    try:
        response = model.generate_content(prompt, request_options={"timeout": 600})
        cleaned_response = re.sub(r'```json\s*|\s*```', '', response.text, flags=re.DOTALL).strip()
        translated_texts = json.loads(cleaned_response)
        if len(translated_texts) != len(transcript_data):
            st.warning("Translation mismatch. Some segments may not be translated.", icon="⚠️")
            return None
        translated_data = [dict(segment, text=translated_texts[i]) for i, segment in enumerate(transcript_data)]
        return translated_data
    except Exception as e:
        st.error(f"Translation failed: {e}", icon="❌")
        return None

def _offset_timestamps(transcription_text, offset_seconds):
    """Helper function to add a time offset to timestamps."""
    def replacer(match):
        h, m, s, ms = map(int, match.groups())
        original_time = timedelta(hours=h, minutes=m, seconds=s, milliseconds=ms)
        offset_delta = timedelta(seconds=offset_seconds)
        new_time = original_time + offset_delta
        total_seconds_val = new_time.total_seconds()
        new_h, new_m, new_s, new_ms = int(total_seconds_val // 3600), int((total_seconds_val % 3600) // 60), int(total_seconds_val % 60), int(new_time.microseconds / 1000)
        return f"[{new_h:02d}:{new_m:02d}:{new_s:02d}.{new_ms:03d}]"
    pattern = re.compile(r'\[(\d{2}):(\d{2}):(\d{2})\.(\d{3})\]')
    return pattern.sub(replacer, transcription_text)

@st.cache_data(show_spinner="Transcribing chunk...", persist=True)
def _transcribe_chunk(audio_chunk_data, model_name):
    """A cached function to transcribe a single audio chunk."""
    tmp_file_path, gemini_file = None, None
    model = genai.GenerativeModel(model_name=model_name)
    prompt = (
        "You are a highly accurate audio transcription service. Your task is to transcribe the provided Bengali audio file with maximum precision. "
        "Follow these rules strictly:\n"
        "1.  **Output Format:** The transcription must be in Bengali script (Unicode).\n"
        "2.  **Speaker Diarization:** Identify each speaker and label them as 'বক্তা ১:', 'বক্তা ২:', etc.\n"
        "3.  **Timestamp Precision:** This is the most critical rule. You MUST provide a precise start timestamp in `[HH:MM:SS.mmm]` format at the beginning of each speaker's turn. The timestamp must be relative to the start of the audio chunk I provide.\n"
        "4.  **Segmentation:** Create a new timestamped line after a significant pause (more than 2-3 seconds).\n"
        "5.  **No Extraneous Text:** Your output should only contain timestamps, speaker labels, and the transcribed text.\n"
    )
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
            tmp_file.write(audio_chunk_data); tmp_file_path = tmp_file.name
        gemini_file = genai.upload_file(path=tmp_file_path)
        response = model.generate_content([prompt, gemini_file], request_options={"timeout": 600})
        if not response.parts:
            reason = "Unknown"
            if response.prompt_feedback and response.prompt_feedback.block_reason: reason = response.prompt_feedback.block_reason.name
            return f"Error: Chunk transcription failed (Reason: {reason})."
        return response.parts[0].text
    finally:
        if tmp_file_path and os.path.exists(tmp_file_path): os.remove(tmp_file_path)
        if gemini_file: genai.delete_file(gemini_file.name)

def transcribe_audio_with_gemini(file_content, file_type, model_name):
    """Main transcription orchestrator with a resilient, 'bulletproof' loop."""
    CHUNK_LENGTH_MINUTES = 10; CHUNK_LENGTH_MS = CHUNK_LENGTH_MINUTES * 60 * 1000; API_CALL_DELAY_SECONDS = 5
    st.info("Step 1: Normalizing audio...")
    audio = AudioSegment.from_file(io.BytesIO(file_content), format=file_type)
    processed_audio = audio.set_frame_rate(16000).set_channels(1).apply_gain(-20.0 - audio.dBFS)
    duration_ms = len(processed_audio)
    if duration_ms <= CHUNK_LENGTH_MS:
        st.info("Audio is short. Transcribing...")
        with io.BytesIO() as audio_io:
            processed_audio.export(audio_io, format="wav")
            return _transcribe_chunk(audio_io.getvalue(), model_name)
    else:
        num_chunks = math.ceil(duration_ms / CHUNK_LENGTH_MS)
        st.info(f"Audio is long. Splitting into {num_chunks} chunks.")
        all_transcriptions = []
        progress_bar = st.progress(0, text="Transcribing chunks...")
        for i in range(num_chunks):
            if i > 0: time.sleep(API_CALL_DELAY_SECONDS)
            start_ms, end_ms = i * CHUNK_LENGTH_MS, (i + 1) * CHUNK_LENGTH_MS
            audio_chunk = processed_audio[start_ms:end_ms]
            st.info(f"Processing chunk {i+1}/{num_chunks}...")
            try:
                with io.BytesIO() as audio_io:
                    audio_chunk.export(audio_io, format="wav")
                    chunk_data = audio_io.getvalue()
                raw_chunk_transcription = _transcribe_chunk(chunk_data, model_name)
                if not raw_chunk_transcription or "Error:" in raw_chunk_transcription:
                    raise ValueError(f"Transcription failed for chunk {i+1}.")
                all_transcriptions.append(_offset_timestamps(raw_chunk_transcription, start_ms / 1000))
            except Exception as e:
                st.error(f"Error on chunk {i+1}: {e}")
                time_obj = timedelta(seconds=start_ms / 1000)
                h, m, s = int(time_obj.total_seconds()//3600), int((time_obj.total_seconds()%3600)//60), int(time_obj.total_seconds()%60)
                all_transcriptions.append(f"[{h:02d}:{m:02d}:{s:02d}.000] বক্তা ?: [TRANSCRIPTION FAILED]")
            progress_bar.progress((i + 1) / num_chunks)
        st.success("All chunks processed!"); return "\n".join(all_transcriptions)

@st.cache_data(show_spinner="Analyzing text...")
def analyze_text_with_gemini(text, task):
    model = genai.GenerativeModel(model_name="models/gemini-1.5-flash-latest")
    prompts = {"summarize": f"Summarize: {text}", "topics": f"List topics in: {text}"}
    return model.generate_content(prompts.get(task, "Unknown task")).text

@st.cache_data(show_spinner="Identifying speakers...")
def identify_speaker_names(transcript_text, speaker_labels):
    model = genai.GenerativeModel(model_name="models/gemini-1.5-flash-latest")
    labels_str = ", ".join(speaker_labels)
    prompt = (f"Analyze transcript. Speakers: {labels_str}. Find names. Reply ONLY with JSON. Keep original labels if unknown. Example: {{\"বক্তা ১:\": \"রহিম:\", \"বক্তা ২:\": \"বক্তা ২:\"}}\n\n{transcript_text}")
    try:
        response = model.generate_content(prompt); cleaned_json_str = re.sub(r'```json\s*|\s*```', '', response.text, flags=re.DOTALL).strip()
        return json.loads(cleaned_json_str)
    except Exception: return {label: label for label in speaker_labels}

def parse_timestamped_transcription(raw_text):
    data = []; pattern = re.compile(r'\[(\d{2}):(\d{2}):(\d{2})\.(\d{3})\]\s*(বক্তা\s*\d+|\?+)\s*:\s*([\s\S]*?)(?=\[\d{2}:\d{2}:\d{2}\.\d{3}\]|\Z)')
    for i, match in enumerate(pattern.finditer(raw_text)):
        h, m, s, ms, speaker, text = match.groups()
        total_seconds = int(h) * 3600 + int(m) * 60 + int(s) + int(ms) / 1000.0
        data.append({"id": f"segment_{i}", "time_sec": total_seconds, "timestamp": f"[{h}:{m}:{s}]", "speaker": speaker.strip() + ":", "text": text.strip()})
    return data

def get_full_transcript_text(transcript_data, speaker_map):
    return "\n\n".join([f"{speaker_map.get(entry['speaker'], entry['speaker'])}\n{entry['text']}" for entry in transcript_data])

def create_docx_content(transcript_data, speaker_map):
    doc = Document(); doc.add_heading('Audio Transcription', 1)
    for entry in transcript_data:
        p = doc.add_paragraph(); p.add_run(f"{speaker_map.get(entry['speaker'], entry['speaker'])}\n").bold = True; p.add_run(entry['text'])
    bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio.getvalue()
    
def create_transcript_display(transcript_data, speaker_map):
    js_transcript_data = json.dumps(transcript_data)
    transcript_html = ""
    for i, entry in enumerate(transcript_data):
        renamed_speaker = speaker_map.get(entry['speaker'], entry['speaker'])
        escaped_text = (entry['text'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;').replace("'","&#39;"))
        transcript_html += f"""<div class="transcript-card" id="segment-{i}"><div class="card-header"><div class="speaker-label">{renamed_speaker}</div><div class="timestamp-btn" onclick="seekAudio({entry['time_sec']})"><svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" fill="currentColor" class="bi bi-play-circle-fill" viewBox="0 0 16 16"><path d="M16 8A8 8 0 1 1 0 8a8 8 0 0 1 16 0zM6.79 5.093A.5.5 0 0 0 6 5.5v5a.5.5 0 0 0 .79.407l3.5-2.5a.5.5 0 0 0 0-.814l-3.5-2.5z"/></svg><span>{entry['timestamp']}</span></div></div><div class="card-body">{escaped_text}</div></div>"""
    
    component_html = f"""<head><link href="https://fonts.googleapis.com/css2?family=Noto+Sans+Bengali:wght@400;700&family=Roboto+Mono:wght@400;700&display=swap" rel="stylesheet"></head><body><style>#transcript-container{{border:1px solid #e0e0e0;border-radius:8px;max-height:60vh;overflow-y:auto;padding:10px;background-color:#f8f9fa}}.transcript-card{{background-color:#fff;border-radius:6px;padding:12px 16px;margin-bottom:12px;border:1px solid #e0e0e0;transition:all .2s ease-in-out}}.transcript-card.highlight{{border-left:5px solid #0d6efd;box-shadow:0 4px 12px rgba(0,0,0,.08);transform:scale(1.01)}}.card-header{{display:flex;justify-content:space-between;align-items:center;margin-bottom:8px}}.speaker-label{{font-family:'Noto Sans Bengali',sans-serif;font-weight:700;color:#0d6efd;font-size:1.1em}}.timestamp-btn{{font-family:'Roboto Mono',monospace;font-size:.9em;color:#495057;background-color:#e9ecef;padding:4px 8px;border-radius:4px;cursor:pointer;display:flex;align-items:center;gap:5px;transition:background-color .2s}}.timestamp-btn:hover{{background-color:#ced4da}}.card-body{{font-family:'Noto Sans Bengali',sans-serif;line-height:1.7;color:#212529;font-size:1.05em;text-align:justify}}</style><div id="transcript-container">{transcript_html}</div><script>const transcriptData={js_transcript_data},audioPlayer=parent.document.querySelector("audio"),transcriptContainer=document.getElementById("transcript-container");let lastHighlightedIndex=-1;function seekAudio(e){{audioPlayer&&(audioPlayer.currentTime=e,audioPlayer.play())}}function highlightCurrentSegment(){{if(audioPlayer){{const e=audioPlayer.currentTime;let t=-1;for(let o=transcriptData.length-1;o>=0;o--)if(e>=transcriptData[o].time_sec){{t=o;break}}if(t!==lastHighlightedIndex){{-1!==lastHighlightedIndex&&document.getElementById(`segment-${{lastHighlightedIndex}}`)?.classList.remove("highlight"),-1!==t&&(currentSegment=document.getElementById(`segment-${{t}}`))?.classList.add("highlight"),"textarea"!==document.activeElement.tagName.toLowerCase()&&currentSegment?.scrollIntoView({{behavior:"smooth",block:"center"}}),lastHighlightedIndex=t}}}}}}audioPlayer?.addEventListener("timeupdate",highlightCurrentSegment);</script></body>"""
    st.components.v1.html(component_html, height=600)

def create_analytics_dashboard(transcript_data, speaker_map):
    st.header("📊 Conversation Dashboard")
    if len(transcript_data) < 2: st.info("Analytics require at least two transcribed segments."); return
    try:
        df = pd.DataFrame(transcript_data); df['next_time_sec'] = df['time_sec'].shift(-1)
        if len(df) > 1: avg_dur = (df['next_time_sec'] - df['time_sec']).mean(); total_dur = df['time_sec'].iloc[-1] + avg_dur; df['next_time_sec'].fillna(total_dur, inplace=True)
        else: df['next_time_sec'].fillna(df['time_sec'] + 5, inplace=True)
        df['duration'] = df['next_time_sec'] - df['time_sec']; df['duration'] = df['duration'].apply(lambda x: max(x, 0)); df['renamed_speaker'] = df['speaker'].map(speaker_map)
        col1, col2 = st.columns(2)
        with col1:
            talk_time = df.groupby('renamed_speaker')['duration'].sum().reset_index()
            fig_pie = px.pie(talk_time, values='duration', names='renamed_speaker', title='<b>Speaker Contribution</b>', hole=0.3, color_discrete_sequence=px.colors.qualitative.Set2)
            fig_pie.update_traces(textposition='inside', textinfo='percent+label', hovertemplate='Speaker: %{label}<br>Talk Time: %{value:.1f}s<br>Contribution: %{percent}')
            fig_pie.update_layout(showlegend=False, title_x=0.5, font=dict(family="Arial, sans-serif")); st.plotly_chart(fig_pie, use_container_width=True)
        with col2:
            fig_timeline = px.bar(df, x='time_sec', y='duration', color='renamed_speaker', title='<b>Conversation Activity Timeline</b>', labels={'time_sec': 'Time (s)', 'duration': 'Duration (s)'}, hover_name='renamed_speaker', hover_data={'text': False})
            fig_timeline.update_layout(xaxis_title=None, yaxis_title="Speech Duration", title_x=0.5, font=dict(family="Arial, sans-serif")); st.plotly_chart(fig_timeline, use_container_width=True)
    except Exception as e: st.warning(f"Could not generate analytics dashboard. Error: {e}")

# --- MAIN APP ---
def main():
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated=False; st.session_state.username=None; st.session_state.user_id=None; st.session_state.role=None; st.session_state.transcript_data=None; st.session_state.speaker_map={}; st.session_state.last_uploaded_file_id=None; st.session_state.current_transcript_title="Untitled Transcript"; st.session_state.original_filename=None; st.session_state.edit_mode=False; st.session_state.selected_model="models/gemini-1.5-flash-latest"; st.session_state.translate_mode=False; st.session_state.translated_data=None

    if not st.session_state.authenticated:
        st.header("Bengali Transcriber Login")
        with st.form("login_form"):
            username = st.text_input("Username"); password = st.text_input("Password", type="password")
            if st.form_submit_button("Login"):
                user = authenticate_user(username, password)
                if user: st.session_state.authenticated = True; st.session_state.username = user.username; st.session_state.user_id = user.id; st.session_state.role = user.role; st.rerun()
                else: st.error("Invalid username or password")
    else:
        st.sidebar.success(f"Logged in as **{st.session_state.username}**")
        if st.sidebar.button("Logout"):
            for key in list(st.session_state.keys()): del st.session_state[key]
            st.rerun()
        st.sidebar.markdown("---"); st.sidebar.header(" My Transcripts")
        user_transcripts = get_user_transcripts(st.session_state.user_id)
        if not user_transcripts: st.sidebar.info("You have no saved transcripts.")
        else:
            for t in user_transcripts:
                if st.sidebar.button(f"📄 {t.title}", key=f"load_{t.id}"):
                    st.session_state.transcript_data = t.content; st.session_state.current_transcript_title = t.title; st.session_state.original_filename = t.original_filename
                    unique_speakers = sorted(list(set(d['speaker'] for d in t.content if 'speaker' in d))); st.session_state.speaker_map = {sp: sp for sp in unique_speakers}
                    st.session_state.translate_mode=False; st.session_state.translated_data=None
                    st.success(f"Loaded '{t.title}'"); st.rerun()
        if st.session_state.get('transcript_data'):
            st.sidebar.markdown("---"); st.sidebar.header("✏️ Edit Speakers")
            st.sidebar.caption("Rename speakers below.")
            for original_speaker in list(st.session_state.speaker_map.keys()):
                new_name = st.sidebar.text_input(f"Rename {original_speaker}", value=st.session_state.speaker_map[original_speaker], key=f"rename_{original_speaker}")
                st.session_state.speaker_map[original_speaker] = new_name
        st.sidebar.markdown("---"); st.sidebar.header("⚙️ Settings")
        model_options = ["models/gemini-1.5-flash-latest", "models/gemini-1.5-pro-latest"]
        st.session_state.selected_model = st.selectbox("Choose Gemini Model", options=model_options, index=model_options.index(st.session_state.selected_model), help="Flash is faster. Pro is more accurate.")
        if st.session_state.role == 'admin':
            with st.sidebar.expander("👑 Admin Panel"):
                st.subheader("Add New User");
                with st.form("add_user_form", clear_on_submit=True):
                    new_user = st.text_input("New Username"); new_pass = st.text_input("New Password", type="password"); new_role = st.selectbox("Role", ["user", "admin"])
                    if st.form_submit_button("Add User"):
                        if new_user and new_pass:
                            success, message = add_new_user(new_user, new_pass, new_role)
                            if success: st.success(message)
                            else: st.error(message)
                st.subheader("All Users"); st.dataframe(get_all_users_from_db(), use_container_width=True)
        st.title("Pro Bengali Transcription Workbench 🎧")
        st.markdown("For audio >10 mins, the app automatically splits them for better processing.")
        uploaded_file = st.file_uploader("1. Upload a new audio file to start", type=["wav", "mp3", "m4a"])
        if uploaded_file:
            if uploaded_file.file_id != st.session_state.get('last_uploaded_file_id'):
                st.info("New audio detected. Clearing session."); st.session_state.transcript_data = None; st.session_state.speaker_map = {}; st.session_state.last_uploaded_file_id = uploaded_file.file_id; st.session_state.original_filename = uploaded_file.name; st.session_state.current_transcript_title = os.path.splitext(uploaded_file.name)[0]; st.cache_data.clear(); st.session_state.translate_mode=False; st.session_state.translated_data=None
            audio_bytes = uploaded_file.getvalue(); st.audio(audio_bytes)
            if st.button("2. Transcribe Audio", type="primary"):
                st.session_state.translate_mode=False; st.session_state.translated_data=None
                raw_format = uploaded_file.type.split('/')[1]; format_map = {'x-m4a': 'm4a', 'mpeg': 'mp3'}; normalized_format = format_map.get(raw_format, raw_format)
                try:
                    with st.spinner("Transcription in progress..."):
                        raw_transcription = transcribe_audio_with_gemini(audio_bytes, normalized_format, st.session_state.selected_model)
                    st.session_state.transcript_data = parse_timestamped_transcription(raw_transcription)
                    if not st.session_state.transcript_data:
                        st.warning("Could not parse timestamps from transcription."); st.text_area("Raw Output", raw_transcription, height=300)
                    else:
                        generic_speakers = sorted(list(set(d['speaker'] for d in st.session_state.transcript_data if d['speaker'] != "বক্তা ?:"))); full_text_for_analysis = get_full_transcript_text(st.session_state.transcript_data, st.session_state.speaker_map); st.session_state.speaker_map = identify_speaker_names(full_text_for_analysis, generic_speakers); st.session_state.edit_mode = False
                        st.success("Transcription complete!")
                except Exception as e:
                    st.error(f"An error occurred: {e}");
                    if "ResourceExhausted" in str(e): st.error("API quota exceeded.")
                    elif "Decoding failed" in str(e) or "Encoding failed" in str(e): st.info("💡 Hint: Audio file might be corrupted.")
        if st.session_state.get('transcript_data'):
            display_data = st.session_state.transcript_data
            create_analytics_dashboard(display_data, st.session_state.speaker_map)
            st.markdown("---"); st.header("3. Review, Edit, and Save")
            st.session_state.translate_mode = st.toggle("Translate to English", value=st.session_state.translate_mode, help="Translate the Bengali transcript to English.")
            st.session_state.edit_mode = st.toggle("Enable Edit Mode", value=st.session_state.edit_mode, help="Switch to a mode where you can correct the text.")
            if st.session_state.translate_mode:
                if st.session_state.translated_data is None:
                    translated_result = _translate_text_with_gemini(st.session_state.transcript_data)
                    if translated_result: st.session_state.translated_data = translated_result
                if st.session_state.translated_data: display_data = st.session_state.translated_data
            if not st.session_state.edit_mode:
                st.caption("Read & Review Mode"); create_transcript_display(display_data, st.session_state.speaker_map)
            else:
                st.caption("Edit Mode")
                with st.form("edit_form"):
                    for i, entry in enumerate(display_data):
                        renamed_speaker = st.session_state.speaker_map.get(entry['speaker'], entry['speaker'])
                        col1, col2 = st.columns([0.15, 0.85]);
                        with col1: st.write(f"**{renamed_speaker}**<br>`{entry['timestamp']}`", unsafe_allow_html=True)
                        with col2: display_data[i]['text'] = st.text_area("Segment text", value=entry['text'], key=f"text_edit_{entry['id']}", label_visibility="collapsed")
                    if st.form_submit_button("✅ Apply Edits & Return to View Mode", type="primary"):
                        if st.session_state.translate_mode: st.session_state.translated_data = display_data
                        else: st.session_state.transcript_data = display_data
                        st.session_state.edit_mode = False; st.success("Edits applied!"); st.rerun()
            st.markdown("---")
            col_save, col_actions = st.columns([0.6, 0.4])
            with col_save:
                with st.form("save_form_2"):
                    title = st.text_input("Transcription Title", value=st.session_state.current_transcript_title)
                    if st.form_submit_button("💾 Save Original Transcript to Database", type="primary"):
                        if save_transcript_to_db(st.session_state.user_id, title, st.session_state.original_filename, st.session_state.transcript_data):
                            st.success(f"Original transcript '{title}' saved!"); st.session_state.current_transcript_title = title
                        else: st.error("Failed to save transcript.")
            with col_actions:
                st.write(" ")
                with st.expander("Analyze & Export", expanded=True):
                    full_text = get_full_transcript_text(display_data, st.session_state.speaker_map)
                    filename_suffix = " (Translated)" if st.session_state.translate_mode else ""
                    export_filename = f"{st.session_state.current_transcript_title}{filename_suffix}"
                    action_cols = st.columns(3)
                    with action_cols[0]:
                        if st.button("Summary"): st.text_area("Summary", analyze_text_with_gemini(full_text, "summarize"), height=150)
                    with action_cols[1]:
                        st.download_button("Export TXT", full_text.encode('utf-8'), f"{export_filename}.txt")
                    with action_cols[2]:
                        st.download_button("Export DOCX", create_docx_content(display_data, st.session_state.speaker_map), f"{export_filename}.docx")

if __name__ == "__main__":
    main()